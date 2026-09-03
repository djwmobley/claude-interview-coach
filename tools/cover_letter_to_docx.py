"""
Convert a plain-text cover letter to a styled DOCX.

Format expected:
  Line 1:        Sender name
  Line 2:        Contact line
  Blank line
  Date line
  Blank line
  Recipient block (one or more lines, ends at blank)
  Blank line
  Re: line (optional: must start with "Re:")
  Blank line
  Body paragraphs (separated by blank lines)
  Blank line
  Closing name

Usage:
  python tools/cover_letter_to_docx.py <input.txt> <output.docx>
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY  = RGBColor(0x0D, 0x21, 0x37)
BLUE  = RGBColor(0x1A, 0x3A, 0x5C)
GREY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

MARGIN_IN  = 1.0
NAME_PT    = 16
CONTACT_PT = 10
META_PT    = 10
BODY_PT    = 11
SPACE_PT   = 6


def set_eastasia_cs_font(style, font_name="Calibri"):
    """python-docx's Font.name setter only sets w:rFonts ascii/hAnsi; without
    eastAsia/cs, a non-Latin character in body text falls back to Word's
    theme font instead of Calibri."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def run(para, text, size=BODY_PT, bold=False, color=BLACK):
    r = para.add_run(text)
    r.font.name      = "Calibri"
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.color.rgb = color
    return r


def para(doc, text="", size=BODY_PT, bold=False, color=BLACK,
         align=None, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if align:
        p.alignment = align
    if text:
        run(p, text, size=size, bold=bold, color=color)
    return p


def parse(lines):
    """Parse cover letter lines into structured sections."""
    # Strip trailing whitespace, keep blank lines as ""
    lines = [l.rstrip() for l in lines]

    # Remove leading/trailing blank lines
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    sections = {"name": "", "contact": "", "date": "", "recipient": [],
                "re": "", "body": [], "closing": ""}

    idx = 0

    # Name
    if idx < len(lines):
        sections["name"] = lines[idx]; idx += 1

    # Contact
    if idx < len(lines) and lines[idx]:
        sections["contact"] = lines[idx]; idx += 1

    # Skip blank
    while idx < len(lines) and not lines[idx]:
        idx += 1

    # Date
    if idx < len(lines):
        sections["date"] = lines[idx]; idx += 1

    # Skip blank
    while idx < len(lines) and not lines[idx]:
        idx += 1

    # Recipient block (lines until blank)
    while idx < len(lines) and lines[idx]:
        sections["recipient"].append(lines[idx]); idx += 1

    # Skip blank
    while idx < len(lines) and not lines[idx]:
        idx += 1

    # Optional Re: line
    if idx < len(lines) and lines[idx].startswith("Re:"):
        sections["re"] = lines[idx]; idx += 1
        while idx < len(lines) and not lines[idx]:
            idx += 1

    # Body paragraphs and closing name
    current = []
    while idx < len(lines):
        line = lines[idx]
        if line:
            current.append(line)
        else:
            if current:
                sections["body"].append(" ".join(current))
                current = []
        idx += 1
    if current:
        sections["body"].append(" ".join(current))

    # Last body item that is just a name (no punctuation, short) = closing
    if sections["body"] and len(sections["body"][-1].split()) <= 4 \
            and not sections["body"][-1].endswith("."):
        sections["closing"] = sections["body"].pop()

    return sections


def convert(txt_path, docx_path):
    with open(txt_path, encoding="utf-8") as f:
        lines = f.readlines()

    s = parse(lines)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin   = Inches(MARGIN_IN)
        section.right_margin  = Inches(MARGIN_IN)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT)
    set_eastasia_cs_font(normal)

    # Name
    para(doc, s["name"], size=NAME_PT, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=1)

    # Contact
    if s["contact"]:
        para(doc, s["contact"], size=CONTACT_PT, color=GREY, after=SPACE_PT)

    # Thin rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Date
    if s["date"]:
        para(doc, s["date"], size=META_PT, color=BLACK, after=SPACE_PT)

    # Recipient
    for line in s["recipient"]:
        p2 = para(doc, line, size=META_PT, color=BLACK, after=0)

    if s["recipient"]:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(SPACE_PT)

    # Re: line
    if s["re"]:
        para(doc, s["re"], size=META_PT, bold=True, color=NAVY, after=SPACE_PT * 2)

    # Body paragraphs
    for i, body_para in enumerate(s["body"]):
        after = SPACE_PT if i < len(s["body"]) - 1 else SPACE_PT * 2
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(after)
        run(p3, body_para, size=BODY_PT, color=BLACK)

    # Closing name
    if s["closing"]:
        para(doc, s["closing"], size=BODY_PT, bold=False, color=BLACK, before=SPACE_PT)

    doc.save(docx_path)
    print(f"Cover letter written to {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python tools/cover_letter_to_docx.py <input.txt> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
