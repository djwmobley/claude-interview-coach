"""
Convert a markdown resume (claude-interview-coach format) to a styled DOCX.

Document structure expected:
  Block 0  | Name / contact / tagline
  ---
  Block 1  | Summary paragraphs        -> rendered under "PROFESSIONAL SUMMARY"
  ---
  Block 2  | Competencies              -> rendered under "CORE COMPETENCIES"
  ---
  Block 3+ | Experience, Education, Certifications

Within experience blocks:
  ALLCAPS          -> section heading  (e.g. EXPERIENCE)
  Title line       -> role title       (next non-blank line contains |)
  Pipe line        -> company line     (contains |)
  Description line -> company context  (follows a pipe line, no |)
  · Bullet         -> experience bullet (continuation lines start with 2+ spaces)

Page count is not tuned to a fixed cap: 2 pages is the target, 3 is
acceptable. Content is never cut to fit a page count. Job blocks never split
across a page boundary; a trailing section with no bullets of its own (e.g.
CERTIFICATIONS right after EDUCATION) stays on the same page as the section
before it, regardless of section name (see PageState / add_section_heading).
"""

import re
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0D, 0x21, 0x37)   # name, section headings
BLUE  = RGBColor(0x1A, 0x3A, 0x5C)   # role titles, tagline
GREY  = RGBColor(0x55, 0x55, 0x55)   # contact line, company descriptions
BLACK = RGBColor(0x1A, 0x1A, 0x1A)   # body text

# ── Size / spacing constants (tune here to control page count) ────────────────
# Every size below NAME_PT is 10pt (a single uniform body size, ATS-safe).
# NAME_PT is the sole exception, used as a title treatment for the name.
NAME_PT         = 18
TAGLINE_PT      = 10
CONTACT_PT      = 10
SECTION_HDR_PT  = 10
ROLE_PT         = 10
BODY_PT         = 10

MARGIN_IN       = 0.65   # all four sides

DIVIDER_SPACE   = 3      # pt before/after thin rule
SUMMARY_SPACE   = 2      # pt before/after summary paragraphs
ROLE_BEFORE     = 7      # pt before role title
DESC_AFTER      = 3      # pt after company description (before first bullet)
BULLET_SPACE    = 0.5    # pt before/after each bullet


# ── Pagination state ──────────────────────────────────────────────────────────

class PageState:
    """Tracks whether the keep-together chain running through section content
    should be released at the next section heading (spec amendment A4: a
    positional rule, not a heading-name allow-list). A section that adds no
    bullets of its own (e.g. CERTIFICATIONS immediately after EDUCATION)
    never releases the chain, so it stays on the same page as the section
    before it -- regardless of what either section is named. A section that
    does add bullets (e.g. EXPERIENCE) always releases the chain at its own
    end, so the next section is free to start a new page.
    """

    def __init__(self):
        self.seen_heading = False
        self.bullets_since_heading = 0


# ── Low-level helpers ─────────────────────────────────────────────────────────

def run(para, text, size=BODY_PT, bold=False, italic=False, color=BLACK):
    r = para.add_run(text)
    r.font.name = "Calibri"
    if size != BODY_PT:
        r.font.size = Pt(size)
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.color.rgb = color
    return r


def spacing(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def keep_with_next(para):
    para.paragraph_format.keep_with_next = True


def keep_together(para):
    para.paragraph_format.keep_together = True


def close_block(doc):
    """Release keep_with_next from the last paragraph, closing the current keep-together chain."""
    if not doc.paragraphs:
        return
    doc.paragraphs[-1].paragraph_format.keep_with_next = None


def bottom_border(para, color="0D2137", size="8"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def thin_rule(para):
    bottom_border(para, color="CCCCCC", size="4")


def strip_contextual_spacing(style):
    """Remove w:contextualSpacing from a style's pPr, if present. Word treats
    contextualSpacing as "collapse spacing between paragraphs of this style",
    which would otherwise override BULLET_SPACE between adjacent bullets."""
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        return
    cs = pPr.find(qn("w:contextualSpacing"))
    if cs is not None:
        pPr.remove(cs)


def set_eastasia_cs_font(style, font_name="Calibri"):
    """python-docx's Font.name setter only sets w:rFonts ascii/hAnsi; without
    eastAsia/cs, a non-Latin character in body text falls back to Word's
    theme font instead of Calibri."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


# ── Pre-processing ────────────────────────────────────────────────────────────

def join_continuations(lines):
    """Merge 2-space-indented continuation lines into their parent bullet."""
    out = []
    for line in lines:
        if line.startswith("  ") and out and out[-1].strip():
            out[-1] = out[-1].rstrip() + " " + line.strip()
        else:
            out.append(line)
    return out


def split_by_dividers(lines):
    blocks, current = [], []
    for line in lines:
        if line.strip() == "---":
            blocks.append(current)
            current = []
        else:
            current.append(line)
    blocks.append(current)
    return blocks


# ── Block renderers ───────────────────────────────────────────────────────────

def add_divider(doc):
    p = doc.add_paragraph()
    spacing(p, before=DIVIDER_SPACE, after=DIVIDER_SPACE)
    thin_rule(p)


def add_section_heading(doc, text, state):
    # A4: only release the previous keep-together chain if at least one
    # bullet was rendered since the previous heading, or there was no
    # previous heading at all. A short trailing section with no bullets of
    # its own (CERTIFICATIONS after EDUCATION) never releases the chain.
    if not state.seen_heading or state.bullets_since_heading > 0:
        close_block(doc)
    state.seen_heading = True
    state.bullets_since_heading = 0

    p = doc.add_paragraph()
    spacing(p, before=8, after=2)
    run(p, text.upper(), size=SECTION_HDR_PT, bold=True, color=NAVY)
    bottom_border(p)
    keep_with_next(p)


def render_header(doc, lines):
    non_blank = [l.rstrip() for l in lines if l.strip()]
    if not non_blank:
        return

    # Name: centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, before=0, after=1)
    run(p, non_blank[0], size=NAME_PT, bold=True, color=NAVY)

    # Contact: centered
    if len(non_blank) > 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=0, after=3)
        run(p, non_blank[1], size=CONTACT_PT, color=GREY)

    # Tagline: centered
    if len(non_blank) > 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=1, after=0)
        run(p, non_blank[2], size=TAGLINE_PT, color=BLUE)


def render_summary(doc, lines, state):
    add_section_heading(doc, "Professional Summary", state)
    para_lines = []
    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            if para_lines:
                p = doc.add_paragraph()
                spacing(p, before=SUMMARY_SPACE, after=SUMMARY_SPACE)
                run(p, " ".join(para_lines), size=BODY_PT)
                para_lines = []
        else:
            para_lines.append(stripped)
    if para_lines:
        p = doc.add_paragraph()
        spacing(p, before=SUMMARY_SPACE, after=SUMMARY_SPACE)
        run(p, " ".join(para_lines), size=BODY_PT)


def render_competencies(doc, lines, state):
    add_section_heading(doc, "Core Competencies", state)
    text = " ".join(l.strip() for l in lines if l.strip())
    if text:
        p = doc.add_paragraph()
        spacing(p, before=2, after=2)
        run(p, text, size=BODY_PT, color=BLACK)


def add_role_title(doc, text):
    close_block(doc)
    p = doc.add_paragraph()
    spacing(p, before=ROLE_BEFORE, after=0)
    keep_with_next(p)
    run(p, text, size=ROLE_PT, bold=True, color=BLUE)


def add_company_line(doc, text):
    parts = [p.strip() for p in text.split("|")]
    p = doc.add_paragraph()
    spacing(p, before=0, after=0)
    keep_with_next(p)
    run(p, parts[0], size=BODY_PT, bold=True, color=NAVY)
    for part in parts[1:]:
        run(p, "  |  " + part, size=BODY_PT, color=BLACK)


def add_company_description(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=0, after=DESC_AFTER)
    keep_with_next(p)
    run(p, text, size=BODY_PT, italic=True, color=GREY)


def add_bullet(doc, text, state):
    p = doc.add_paragraph(style="List Bullet")
    spacing(p, before=BULLET_SPACE, after=BULLET_SPACE)
    keep_together(p)
    keep_with_next(p)   # chains bullet into block; closed by next role/heading
    run(p, text.lstrip("·").strip(), size=BODY_PT)
    state.bullets_since_heading += 1


def add_plain(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=1, after=1)
    run(p, text, size=BODY_PT)


def render_body(doc, lines, state):
    indexed = [(i, lines[i].rstrip()) for i in range(len(lines))]

    def next_non_blank(i):
        j = i + 1
        while j < len(indexed) and not indexed[j][1].strip():
            j += 1
        return indexed[j][1] if j < len(indexed) else ""

    def prev_non_blank(i):
        k = i - 1
        while k >= 0 and not indexed[k][1].strip():
            k -= 1
        return indexed[k][1] if k >= 0 else ""

    for i, line in indexed:
        text = line.strip()
        if not text:
            continue

        # Section heading: ALL CAPS, letters/spaces only
        if text == text.upper() and re.fullmatch(r"[A-Z\s]+", text) and len(text) > 3:
            add_section_heading(doc, text, state)
            continue

        # Bullet
        if text.startswith("·"):
            add_bullet(doc, text, state)
            continue

        # Company / pipe line
        if "|" in text:
            add_company_line(doc, text)
            continue

        # Role title: no pipe, but next non-blank has pipe
        if "|" in next_non_blank(i):
            add_role_title(doc, text)
            continue

        # Company description: previous non-blank had pipe
        if "|" in prev_non_blank(i):
            add_company_description(doc, text)
            continue

        # Plain text
        add_plain(doc, text)


# ── Entry point ───────────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    doc = Document()
    state = PageState()

    for section in doc.sections:
        section.top_margin    = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin   = Inches(MARGIN_IN)
        section.right_margin  = Inches(MARGIN_IN)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT)
    set_eastasia_cs_font(normal)

    strip_contextual_spacing(doc.styles["List Bullet"])

    with open(md_path, encoding="utf-8") as f:
        raw = f.readlines()

    lines  = [l.rstrip("\n") for l in raw]
    lines  = join_continuations(lines)
    blocks = split_by_dividers(lines)

    if len(blocks) > 0:
        render_header(doc, blocks[0])
    add_divider(doc)

    if len(blocks) > 1:
        render_summary(doc, blocks[1], state)
    add_divider(doc)

    if len(blocks) > 2:
        render_competencies(doc, blocks[2], state)
    add_divider(doc)

    for block in blocks[3:]:
        render_body(doc, block, state)

    close_block(doc)
    doc.save(docx_path)
    print(f"DOCX written to {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python tools/md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    md_file   = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.rsplit(".", 1)[0] + ".docx"
    convert(md_file, docx_file)
