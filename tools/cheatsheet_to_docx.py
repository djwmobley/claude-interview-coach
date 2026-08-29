"""
Convert a markdown call cheat sheet to a styled DOCX.

Supported markdown:
  # Title          -> document title (large, navy)
  ## Section       -> section heading (navy, bold, with rule below)
  ### Subsection   -> subsection heading (blue, bold)
  - bullet         -> indented bullet
  **text**         -> bold inline (within paragraphs and bullets)
  ---              -> thin horizontal rule
  blank line       -> paragraph spacing
  everything else  -> body paragraph

Usage:
  python tools/cheatsheet_to_docx.py <input.md> <output.docx>
"""

import re
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0D, 0x21, 0x37)
BLUE  = RGBColor(0x1A, 0x3A, 0x5C)
GREY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

# ── Sizes ─────────────────────────────────────────────────────────────────────
TITLE_PT      = 16
H2_PT         = 11
H3_PT         = 10
BODY_PT       = 10
BULLET_PT     = 10
SMALL_PT      = 9

MARGIN_IN     = 0.85


# ── Helpers ───────────────────────────────────────────────────────────────────

def add_run(para, text, size=BODY_PT, bold=False, color=BLACK):
    r = para.add_run(text)
    r.font.name      = "Calibri"
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.color.rgb = color
    return r


def add_inline(para, text, size=BODY_PT, color=BLACK):
    """Render text with **bold** spans inline."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            add_run(para, part[2:-2], size=size, bold=True, color=color)
        elif part:
            add_run(para, part, size=size, bold=False, color=color)


def strip_contextual_spacing(style):
    """Remove w:contextualSpacing from a style's pPr, if present, so spacing
    between adjacent bullets is not collapsed."""
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        return
    cs = pPr.find(qn("w:contextualSpacing"))
    if cs is not None:
        pPr.remove(cs)


def set_eastasia_cs_font(style, font_name="Calibri"):
    """python-docx's Font.name setter only sets w:rFonts ascii/hAnsi; without
    eastAsia/cs, a non-Latin character in body text falls back to Word's
    theme font instead of Calibri."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_rule(doc, color_hex="CCCCCC", space_before=4, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, size=TITLE_PT, bold=True, color=NAVY)
    add_rule(doc, color_hex="1A3A5C", space_before=2, space_after=8)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text.upper(), size=H2_PT, bold=True, color=NAVY)
    add_rule(doc, color_hex="CCCCCC", space_before=1, space_after=4)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, text, size=H3_PT, bold=True, color=BLUE)


def add_body(doc, text, before=1, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    add_inline(p, text, size=BODY_PT, color=BLACK)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    add_inline(p, text, size=BULLET_PT, color=BLACK)


# ── Parser ────────────────────────────────────────────────────────────────────

def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = [l.rstrip() for l in f.readlines()]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin   = Inches(MARGIN_IN)
        section.right_margin  = Inches(MARGIN_IN)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT)
    set_eastasia_cs_font(normal)

    strip_contextual_spacing(doc.styles["List Bullet"])

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            add_h1(doc, line[2:].strip())

        elif line.startswith("## "):
            add_h2(doc, line[3:].strip())

        elif line.startswith("### "):
            add_h3(doc, line[4:].strip())

        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())

        elif line == "---":
            add_rule(doc, color_hex="CCCCCC", space_before=6, space_after=6)

        elif line == "":
            pass  # blank lines handled by spacing on adjacent elements

        else:
            add_body(doc, line)

        i += 1

    doc.save(docx_path)
    print(f"Cheat sheet written to {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python tools/cheatsheet_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
